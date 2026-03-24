#!/usr/bin/env python3
"""Generate VenomsBase architecture response docx from markdown source."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path
import re


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    return table


def build_docx(output_path):
    doc = Document()

    # Adjust default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Title
    title = doc.add_heading("VenomsBase: Data Architecture Proposal", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'Response to: "Species-Centered Database Architecture for VenomsBase" '
        "(Castoe, March 2026)"
    )
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ivan Koludarov \u2014 March 2026")
    run.font.size = Pt(10)

    doc.add_paragraph()

    # --- Summary ---
    add_heading(doc, "Summary", 1)
    add_para(
        doc,
        "This document proposes a concrete data architecture for VenomZone, the web platform "
        "for VenomsBase. The central idea: the protein is the primary data object. Each protein "
        "carries its own embedding, sequences, annotations, and links to external data. Genomes "
        "are parallel first-class objects with their own viewer. All other data types \u2014 "
        "transcriptomics, proteomics, structural, functional \u2014 attach as features on "
        "proteins rather than living in separate relational tables.",
    )
    add_para(
        doc,
        "Two working prototypes are attached (interactive HTML viewers for snake and bee venom "
        "proteins) to demonstrate how the protein-centric model works in practice.",
    )

    # --- 1. Protein as Data Container ---
    add_heading(doc, "1. Core Concept: Protein as Data Container", 1)
    add_para(
        doc,
        "A protein language model (ProtT5) encodes each protein sequence as a 1024-dimensional "
        "vector. Proteins with similar sequences, structures, and functions land near each other "
        "in this space \u2014 regardless of species, naming convention, or database of origin. "
        "This gives us a continuous similarity landscape that no relational schema can provide.",
    )
    add_para(doc, "Each protein is a self-contained record with typed data slots:", bold=True)

    # Protein record diagram as a table
    add_table(
        doc,
        ["Slot", "Contents"],
        [
            ["identity", "Accession, gene name, aliases, cross-references (UniProt, GenBank, PDB)"],
            ["sequences", "Full-length (canonical), mature peptide, isoforms / splice variants"],
            ["embeddings", "ProtT5 (1024-dim), ESM2, ESM-C (optional additional models)"],
            ["structure", "AlphaFold prediction, experimental PDB structures"],
            ["taxonomy", "Species, TaxID, common name, clade, order, family"],
            [
                "classification",
                "Venom family + evidence level, venom group/subgroup, "
                "historical classifications (e.g., Dowell 2016)",
            ],
            ["expression", "Tissue-specific TPM/FPKM values, sample metadata"],
            ["proteomics", "Detection evidence (PSMs, confidence), PTMs, proteoforms"],
            ["activity", "Molecular targets, IC50/EC50 values, assay metadata"],
            ["genomic context", "Link to Genome Record (gene neighborhood)"],
            ["literature", "DOIs, PubMed IDs"],
        ],
    )
    doc.add_paragraph()
    add_para(
        doc,
        "New data types (e.g., spatial transcriptomics, cryo-EM) become new slots on the "
        "protein record \u2014 no schema migration required.",
    )

    # --- 2. Three Data Object Types ---
    add_heading(doc, "2. Three Data Object Types", 1)
    add_para(
        doc,
        "The architecture uses three object types, not one hierarchy:",
    )
    add_table(
        doc,
        ["Object", "Role", "Viewer"],
        [
            [
                "Protein Record",
                "Primary data container. Carries embedding, sequences, all annotations.",
                "Embedding scatter (UMAP/PCA) + info panel",
            ],
            [
                "Genome Record",
                "Standalone. Gene annotations, scaffolds, regulatory elements. "
                "Linked to proteins via gene IDs.",
                "Genomic context viewer (gene neighborhood canvas)",
            ],
            [
                "Raw Data",
                "Stored separately, referenced by protein/genome records. "
                "Sequencing reads, mass spectra, expression matrices.",
                "Not directly visualized \u2014 features extracted into protein/genome records",
            ],
        ],
    )
    doc.add_paragraph()
    add_para(
        doc,
        "The key design choice: genomes and proteins are parallel entities, not hierarchical. "
        "A genome region contains multiple genes; each gene may correspond to a protein in the "
        "embedding space. They link to each other, but neither is subordinate.",
    )
    add_para(
        doc,
        "This differs from the species-centered hierarchy (Species \u2192 Genome \u2192 Gene "
        "\u2192 Transcript \u2192 Protein) in an important way: the protein is independently "
        "queryable. You can find similar proteins across all species by nearest-neighbor search "
        "in embedding space, without traversing a join chain.",
    )

    # --- 3. How -Omics Data Fits ---
    add_heading(doc, "3. How -Omics Data Fits", 1)
    add_para(
        doc,
        "All -omics data types follow the same pattern: raw data stored as files, extracted "
        "features become slots on protein records.",
    )
    add_table(
        doc,
        ["Data Type", "Raw Storage", "Feature on Protein", "How It\u2019s Used"],
        [
            [
                "Transcriptomics",
                "FASTQ/BAM files",
                "expression.* (TPM, tissue, stage)",
                "Color by expression; filter by tissue",
            ],
            [
                "Proteomics",
                "MS raw files",
                "proteomics.* (PSMs, PTMs)",
                "Evidence tags; proteoform annotations",
            ],
            [
                "Structure",
                "PDB/mmCIF files",
                "structure.* (AF confidence, method)",
                "3D viewer; structural clustering",
            ],
            [
                "Functional assays",
                "Assay reports",
                "activity.* (IC50, target, method)",
                "Activity overlay; target enrichment",
            ],
            [
                "Genomics",
                "GFF3 + FASTA",
                "genomic_context \u2192 Genome Record link",
                "Gene neighborhood viewer",
            ],
            ["Literature", "\u2014", "literature.* (DOIs)", "Cross-references"],
        ],
    )
    doc.add_paragraph()

    add_para(doc, "Every module in the species-centered schema has a concrete home:", bold=True)
    add_table(
        doc,
        ["Species-Centered Module", "Protein-Centric Equivalent"],
        [
            ["Species & Taxonomy (2A)", "taxonomy.* slots on protein record"],
            ["Genome Assembly (2B)", "Genome Record (standalone object)"],
            ["Gene\u2013Transcript\u2013Protein (2C)", "identity.* + sequences.* slots"],
            ["Venom Classification (2C)", "classification.* slots with evidence level"],
            ["Functional Genomics (3)", "Genome Record regulatory features"],
            ["Transcriptomics (4)", "expression.* slots"],
            ["Proteomics (5)", "proteomics.* slots"],
            ["Structure (6)", "structure.* slots"],
            ["Gene Family (7)", "classification.* slots"],
            ["Functional Activity (8)", "activity.* slots"],
        ],
    )

    # --- 4. Reclassification ---
    add_heading(doc, "4. Why This Matters: The Reclassification Problem", 1)
    add_para(
        doc,
        "Dowell et al. (2016) proposed a venom protein family classification that is widely "
        "used but increasingly recognized as inaccurate. Our studies (Koludarov et al. 2023) "
        "provide updated classifications, but the field still references Dowell\u2019s scheme.",
    )
    p = add_para(doc, "In a relational schema: ", bold=True)
    add_para(
        doc,
        "Updating family assignments means modifying GeneFamily and FamilyMembership tables, "
        "cascading through OrthogroupMembership, and potentially breaking saved queries.",
    )
    p = add_para(doc, "In the protein-centric model: ", bold=True)
    add_para(
        doc,
        "Update the classification.venom_family slot. The old value is preserved as "
        "classification.venom_family_dowell2016. No schema changes, no cascading joins, "
        "both classifications coexist on the same record.",
    )

    # --- 5. Prototypes ---
    add_heading(doc, "5. What the Prototypes Show", 1)
    add_para(doc, "The attached HTML viewers demonstrate three interface components:")
    for label, desc in [
        ("Protein Space", "ProtT5 embedding scatter (UMAP/PCA) as the navigational anchor."),
        ("Information Panel", "All slots displayed on click \u2014 species, family, group, source."),
        (
            "Genomic Context",
            "Gene neighborhood canvas showing venom genes and flanking genes "
            "from published GFF annotations, matched to the clicked protein.",
        ),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(desc)
        run.font.size = Pt(11)

    doc.add_paragraph()
    add_para(
        doc,
        "Snake viewer (2,235 proteins): Pla2g2 (452, Koludarov et al. 2020), KLK serine "
        "proteases (491, Barua & Koludarov 2021 BMC Biology), 3FTx (1,292 toxin proteins, "
        "Koludarov et al. 2023 Nature Comms). Genomic context from published manual GFF annotations.",
    )
    add_para(
        doc,
        "Bee viewer (4,606 proteins): 33 protein families from the Hymenoptera venom arsenal "
        "(Koludarov et al. 2023 BMC Biology). Genomic context from 739 manually annotated GFF3 "
        "regions across 23 species.",
    )

    # --- 6. Tiers ---
    add_heading(doc, "6. Data Completeness Tiers", 1)
    add_table(
        doc,
        ["Tier", "Required Slots", "Example"],
        [
            [
                "Platinum",
                "embedding + sequence + taxonomy + structure + expression + proteomics + activity",
                "Apamin (Apis mellifera)",
            ],
            [
                "Gold",
                "embedding + sequence + taxonomy + structure + classification",
                "Most ToxProt entries",
            ],
            [
                "Silver",
                "embedding + sequence + taxonomy",
                "Computationally predicted venom proteins",
            ],
            ["Bronze", "sequence + taxonomy only", "Unannotated transcriptome hits"],
        ],
    )

    # --- 7. Implementation ---
    add_heading(doc, "7. Implementation Path", 1)
    for phase, desc in [
        (
            "Phase 1 \u2014 Protein records + embedding viewer.",
            "Load existing datasets (ToxProt, our published data) as protein records with "
            "embedding, sequence, taxonomy, and classification slots. Web interface: embedding "
            "scatter with color-by and info panel. This is what the prototypes already demonstrate.",
        ),
        (
            "Phase 2 \u2014 Genome viewer + links.",
            "Add genome records from GFF3 annotations. Link proteins to their genomic context. "
            "Interactive gene neighborhood canvas (already prototyped).",
        ),
        (
            "Phase 3 \u2014 Expression + proteomics + activity.",
            "Add quantitative data slots. Import from published supplementary tables, VenomZone, "
            "and community submissions.",
        ),
        (
            "Phase 4 \u2014 Community features.",
            "User accounts, protein-level comments and annotations, evidence curation, "
            "data submission workflows.",
        ),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(phase + " ")
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(desc)
        run.font.size = Pt(11)

    # --- 8. Next Steps ---
    add_heading(doc, "8. Next Steps", 1)
    for step in [
        "Review the attached viewers and this data model",
        "Agree on the protein record format and slot ontology",
        "Define the initial scope: which species, which data types, what tier to target first",
        "Discuss infrastructure: hosting, data submission pipeline, community curation model",
    ]:
        doc.add_paragraph(step, style="List Bullet")

    doc.add_paragraph()
    add_para(doc, "Happy to jump on a call to walk through it.")

    doc.save(str(output_path))
    size_kb = output_path.stat().st_size / 1024
    print(f"Written: {output_path} ({size_kb:.0f} KB)")


if __name__ == "__main__":
    out = Path(__file__).parent / "VenomsBase_Architecture_Response_Koludarov_2026.docx"
    build_docx(out)
